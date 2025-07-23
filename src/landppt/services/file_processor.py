"""
File Processing Service for LandPPT
Handles document upload and content extraction as specified in requires.md
"""

import os
import re
import logging
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import tempfile

# Document processing libraries
try:
    import docx
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

from ..api.models import FileUploadResponse

logger = logging.getLogger(__name__)


class FileProcessor:
    """Processes uploaded files and extracts content for PPT generation"""
    
    def __init__(self):
        self.supported_formats = {
            '.docx': self._process_docx,
            '.pdf': self._process_pdf,
            '.txt': self._process_txt,
            '.md': self._process_markdown,
            '.jpg': self._process_image,
            '.jpeg': self._process_image,
            '.png': self._process_image,
        }
        
        # Keywords for scenario detection
        self.scenario_keywords = {
            'tourism': ['旅游', '景点', '行程', '旅行', '观光', '度假', '酒店', '机票', '导游'],
            'education': ['教育', '学习', '课程', '培训', '知识', '科普', '儿童', '学生', '教学'],
            'analysis': ['分析', '数据', '统计', '研究', '报告', '调查', '图表', '趋势', '结论'],
            'history': ['历史', '古代', '文化', '传统', '遗产', '文物', '朝代', '事件', '人物'],
            'technology': ['技术', '科技', '创新', '数字', '智能', '人工智能', '互联网', '软件', '硬件'],
            'business': ['商业', '企业', '市场', '营销', '销售', '管理', '战略', '财务', '投资'],
            'general': ['介绍', '概述', '总结', '说明', '展示', '汇报', '演示', '分享']
        }
    
    async def process_file(self, file_path: str, filename: str) -> FileUploadResponse:
        """Process uploaded file and extract content"""
        try:
            file_ext = Path(filename).suffix.lower()
            file_size = os.path.getsize(file_path)
            
            if file_ext not in self.supported_formats:
                raise ValueError(f"Unsupported file format: {file_ext}")
            
            # Process file based on type
            processor = self.supported_formats[file_ext]
            content = await processor(file_path)
            
            # Extract topics and suggest scenarios
            topics = self._extract_topics(content)
            scenarios = self._suggest_scenarios(content)
            
            return FileUploadResponse(
                filename=filename,
                size=file_size,
                type=file_ext,
                processed_content=content,
                extracted_topics=topics,
                suggested_scenarios=scenarios,
                message=f"文件 {filename} 处理成功，提取了 {len(content)} 个字符的内容"
            )
            
        except Exception as e:
            logger.error(f"Error processing file {filename}: {e}")
            raise ValueError(f"文件处理失败: {str(e)}")
    
    async def _process_docx(self, file_path: str) -> str:
        """Process DOCX file"""
        if not DOCX_AVAILABLE:
            raise ValueError("DOCX processing not available. Please install python-docx.")

        def _process_docx_sync(file_path: str) -> str:
            """同步处理DOCX文件（在线程池中运行）"""
            doc = Document(file_path)
            content_parts = []

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    content_parts.append(text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        content_parts.append(" | ".join(row_text))

            return "\n\n".join(content_parts)

        try:
            # 在线程池中执行文件处理以避免阻塞主服务
            import asyncio
            loop = asyncio.get_running_loop()
            return await loop.run_in_executor(None, _process_docx_sync, file_path)

        except Exception as e:
            logger.error(f"Error processing DOCX file: {e}")
            raise ValueError(f"DOCX 文件处理失败: {str(e)}")
    
    async def _process_pdf(self, file_path: str) -> str:
        """Process PDF file"""
        if not PDF_AVAILABLE:
            raise ValueError("PDF processing not available. Please install PyPDF2.")

        def _process_pdf_sync(file_path: str) -> str:
            """同步处理PDF文件（在线程池中运行）"""
            content_parts = []

            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)

                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()

                    if text.strip():
                        content_parts.append(text.strip())

            return "\n\n".join(content_parts)

        try:
            # 在线程池中执行文件处理以避免阻塞主服务
            import asyncio
            loop = asyncio.get_running_loop()
            return await loop.run_in_executor(None, _process_pdf_sync, file_path)

        except Exception as e:
            logger.error(f"Error processing PDF file: {e}")
            raise ValueError(f"PDF 文件处理失败: {str(e)}")
    
    async def _process_txt(self, file_path: str) -> str:
        """Process TXT file"""
        def _process_txt_sync(file_path: str) -> str:
            """同步处理TXT文件（在线程池中运行）"""
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    content = file.read()

                # Try different encodings if UTF-8 fails
                if not content.strip():
                    encodings = ['gbk', 'gb2312', 'latin1']
                    for encoding in encodings:
                        try:
                            with open(file_path, 'r', encoding=encoding) as file:
                                content = file.read()
                            if content.strip():
                                break
                        except:
                            continue

                return content.strip()
            except Exception as e:
                raise e

        try:
            # 在线程池中执行文件处理以避免阻塞主服务
            import asyncio
            loop = asyncio.get_running_loop()
            return await loop.run_in_executor(None, _process_txt_sync, file_path)

        except Exception as e:
            logger.error(f"Error processing TXT file: {e}")
            raise ValueError(f"TXT 文件处理失败: {str(e)}")
    
    async def _process_markdown(self, file_path: str) -> str:
        """Process Markdown file"""
        def _process_markdown_sync(file_path: str) -> str:
            """同步处理Markdown文件（在线程池中运行）"""
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()

            # Remove markdown syntax for cleaner content
            content = re.sub(r'#{1,6}\s+', '', content)  # Remove headers
            content = re.sub(r'\*\*(.*?)\*\*', r'\1', content)  # Remove bold
            content = re.sub(r'\*(.*?)\*', r'\1', content)  # Remove italic
            content = re.sub(r'`(.*?)`', r'\1', content)  # Remove code
            content = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', content)  # Remove links

            return content.strip()

        try:
            # 在线程池中执行文件处理以避免阻塞主服务
            import asyncio
            loop = asyncio.get_running_loop()
            return await loop.run_in_executor(None, _process_markdown_sync, file_path)

        except Exception as e:
            logger.error(f"Error processing Markdown file: {e}")
            raise ValueError(f"Markdown 文件处理失败: {str(e)}")
    
    async def _process_image(self, file_path: str) -> str:
        """Process image file using OCR"""
        if not OCR_AVAILABLE:
            return "图片文件已上传，但 OCR 功能不可用。请安装 pytesseract 和 PIL 以启用文字识别。"

        def _process_image_sync(file_path: str) -> str:
            """同步处理图像文件（在线程池中运行）"""
            image = Image.open(file_path)

            # Perform OCR
            text = pytesseract.image_to_string(image, lang='chi_sim+eng')

            if not text.strip():
                return "图片文件已处理，但未能识别出文字内容。"

            return text.strip()

        try:
            # 在线程池中执行图像处理以避免阻塞主服务
            import asyncio
            loop = asyncio.get_running_loop()
            return await loop.run_in_executor(None, _process_image_sync, file_path)

        except Exception as e:
            logger.error(f"Error processing image file: {e}")
            return f"图片处理失败: {str(e)}"
    
    def _extract_topics(self, content: str) -> List[str]:
        """Extract potential topics from content"""
        if not content:
            return []
        
        topics = []
        
        # Extract sentences that might be topics (short, descriptive)
        sentences = re.split(r'[。！？\n]', content)
        
        for sentence in sentences:
            sentence = sentence.strip()
            # Look for topic-like sentences (10-50 characters, no common words)
            if 10 <= len(sentence) <= 50:
                # Avoid sentences with too many common words
                common_words = ['的', '是', '在', '有', '和', '与', '或', '但', '而', '了', '着', '过']
                common_count = sum(1 for word in common_words if word in sentence)
                
                if common_count <= 2:  # Not too many common words
                    topics.append(sentence)
        
        # Also extract potential titles (lines that are short and at the beginning)
        lines = content.split('\n')
        for i, line in enumerate(lines[:10]):  # Check first 10 lines
            line = line.strip()
            if 5 <= len(line) <= 30 and not line.endswith('：'):
                topics.append(line)
        
        # Remove duplicates and limit to top 10
        topics = list(dict.fromkeys(topics))[:10]
        
        return topics
    
    def _suggest_scenarios(self, content: str) -> List[str]:
        """Suggest appropriate scenarios based on content"""
        if not content:
            return ['general']
        
        content_lower = content.lower()
        scenario_scores = {}
        
        # Score each scenario based on keyword matches
        for scenario, keywords in self.scenario_keywords.items():
            score = 0
            for keyword in keywords:
                score += content_lower.count(keyword)
            
            if score > 0:
                scenario_scores[scenario] = score
        
        # Sort by score and return top scenarios
        sorted_scenarios = sorted(scenario_scores.items(), key=lambda x: x[1], reverse=True)
        
        # Return top 3 scenarios, or 'general' if no matches
        if sorted_scenarios:
            return [scenario for scenario, score in sorted_scenarios[:3]]
        else:
            return ['general']
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats"""
        return list(self.supported_formats.keys())
    
    def validate_file(self, filename: str, file_size: int, max_size_mb: int = 100) -> Tuple[bool, str]:
        """Validate uploaded file"""
        file_ext = Path(filename).suffix.lower()
        
        # Check file extension
        if file_ext not in self.supported_formats:
            return False, f"不支持的文件格式: {file_ext}。支持的格式: {', '.join(self.supported_formats.keys())}"
        
        # Check file size
        max_size_bytes = max_size_mb * 1024 * 1024
        if file_size > max_size_bytes:
            return False, f"文件大小超过限制 ({max_size_mb}MB)。当前文件大小: {file_size / 1024 / 1024:.1f}MB"
        
        # Check specific format requirements
        if file_ext == '.docx' and not DOCX_AVAILABLE:
            return False, "DOCX 处理功能不可用，请联系管理员安装 python-docx"
        
        if file_ext == '.pdf' and not PDF_AVAILABLE:
            return False, "PDF 处理功能不可用，请联系管理员安装 PyPDF2"
        
        if file_ext in ['.jpg', '.jpeg', '.png'] and not OCR_AVAILABLE:
            return True, "图片文件可以上传，但文字识别功能不可用"
        
        return True, "文件验证通过"
    
    async def create_ppt_from_content(self, content: str, suggested_topic: str = None) -> Dict[str, Any]:
        """Create PPT generation request from processed content"""
        # Extract or suggest a topic
        if not suggested_topic:
            topics = self._extract_topics(content)
            suggested_topic = topics[0] if topics else "文档内容展示"
        
        # Suggest scenarios
        scenarios = self._suggest_scenarios(content)
        primary_scenario = scenarios[0] if scenarios else 'general'
        
        # Create a structured outline from content
        sections = self._create_content_sections(content)
        
        return {
            'topic': suggested_topic,
            'scenario': primary_scenario,
            'requirements': f"基于上传文档内容生成PPT，包含以下要点：\n{content[:500]}...",
            'uploaded_content': content,
            'suggested_sections': sections,
            'language': 'zh'
        }
    
    def _create_content_sections(self, content: str) -> List[Dict[str, str]]:
        """Create structured sections from content"""
        sections = []
        
        # Split content into logical sections
        paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
        
        # Create title slide
        sections.append({
            'type': 'title',
            'title': '文档内容展示',
            'subtitle': '基于上传文档生成'
        })
        
        # Create content slides (max 10)
        for i, paragraph in enumerate(paragraphs[:9]):
            if len(paragraph) > 50:  # Only use substantial paragraphs
                # Try to extract a title from the first sentence
                sentences = paragraph.split('。')
                title = sentences[0][:30] + '...' if len(sentences[0]) > 30 else sentences[0]
                
                sections.append({
                    'type': 'content',
                    'title': title or f'内容 {i+1}',
                    'content': paragraph[:300] + '...' if len(paragraph) > 300 else paragraph
                })
        
        # Add thank you slide
        sections.append({
            'type': 'thankyou',
            'title': '谢谢观看',
            'subtitle': '基于文档内容生成'
        })
        
        return sections
